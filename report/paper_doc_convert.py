import json
from openai import  OpenAI
from docx import Document
from pathlib import Path
from utils.content_utils import read_content;
from utils.content_utils import write_str_to_file
client = OpenAI(base_url='https://api.openai-proxy.org/v1', api_key='sk-p8KW4EtRdh7i2MWf9o7YmmQZihySS5HA5D0Z1iEdddtLURpJ');


def generate_prompt_for_check_summary(header,data, table):
    """
    Generate the English prompt used to summarize discrepancies from the `check` field
    across included studies (Table 2, Table 3, Table 4).
    """

    prompt = f"""
You are an expert in systematic review quality verification and data extraction auditing.

Your task is to analyze the provided JSON data, especially the contents of the **"check"** field for each included study, and generate a structured discrepancy summary.

header is {header},
data is {data}
table ttile is {table}

## Dynamic Table Handling
- The dataset may contain discrepancies related to *Table 1, Table 2, Table 3, Table 4,* or other tables.
- **Do NOT assume fixed table numbers.**
- Instead, detect which tables are referenced in the "check" text and generate corresponding sections.
- For every table name that appears (e.g., “Table 2”, “Table 3”, “Table 4”), create a section titled:

### Table X – Main Issues

## Output Requirements
For each detected Table:
- Identify and group recurring discrepancy types into **bullet-point categories**.
- Each issue category must include:
  1. **A short category title**
  2. **A concise explanation**
  3. **Specific study examples** extracted from the `check` field  
     (e.g., “Hamilton et al. [21]”)

## Examples of Possible Issue Categories (not exhaustive):
- Gene panel mismatch or incomplete reporting  
- Baseline demographic inconsistencies  
- Errors in outcome data type or statistical reporting  
- Missing extraction of key variables  
- Incorrect sample size or participant flow  
- Misalignment between reported dataset and actual analysis population  

## Content Rules
- Only use information found in the `check` field of each study.
- Summaries should focus on **issue types**, not rewrite long texts.
- Group similar discrepancies across studies under the same category.
- Maintain an analytical, concise tone.

## Final Output Format Example (structure only)

### Table X – Main Issues
  1.Issue category 1
    brief explanation
    
  2. Issue category 2
     explanation 
- ...


"""

    return prompt


def json_files_to_word(json_files, output_docx):
    """
    将多个 JSON 文件生成 Word 表格，写入同一个文档。

    参数：
        json_files (list of str): JSON 文件路径列表
        output_docx (str): 输出 Word 文件路径
    """
    doc = Document()

    for json_file in json_files:
        # 读取 JSON 文件
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)

        title = data.get('title', 'Table')
        headers = data.get('headers', [])
        table_data = data.get('data', [])

        # 添加标题
        doc.add_heading(title, level=2)

        if not headers or not table_data:
            continue  # 空表格跳过

        # 创建表格
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # 添加表头
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        # 添加数据行
        for row in table_data:
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)

        # 添加空行分隔不同表格
        doc.add_paragraph('')

    # 保存 Word 文件
    doc.save(output_docx)
    print(f"Word 文档已生成: {output_docx}")


# 使用示例
json_folder = Path('D:\\project\\zky\\paperAgent\\check_data')  # JSON 文件存放目录
json_files = list(json_folder.glob('*.json'))  # 获取所有 JSON 文件
results = '';
for json_file in json_files:
    content = read_content(json_file);
    data_content = json.loads(content);
    title = data_content['title'];
    data = data_content['data'];
    data_checks = [];
    for row in data:
        check = row[-1];
        data_checks.append(row);
        #print(f" check is {check}");
    #print(f" content is {content} ");
    print(f" title is {title}");
    headers = data_content['headers'];
    prompt = generate_prompt_for_check_summary(headers, data_checks, title);
    print(f" prompt is {prompt} ");

    response = client.chat.completions.create(
        model="gpt-5",
        messages=[
            {"role": "system", "content": "你是一个医学数据分析专家"},
            {"role": "user", "content": prompt}
        ]
    )
    response_text = response.choices[0].message.content
    results = results + response_text + "\n";


#write_str_to_file("D:\\project\\zky\\paperAgent\\result\\report_3_2_3.txt", results);
#print(f" response_text is {response_text} ");
#print(f" data_check is {data_checks}, title is {title} ");
#print(f" json_files is {json_files} ");
json_files_to_word([str(f) for f in json_files], 'all_studies.docx')
