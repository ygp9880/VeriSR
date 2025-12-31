from docx import Document


def json_to_docx_table(json_data, output_file="output.docx"):
    """
    将 JSON 数据转换为 Word 表格，每行一条研究，error_fields 作为单独列。

    参数:
        json_data: list of dicts, 每个 dict 包含 "study" 和 "error_fields"
        output_file: 输出的 docx 文件名
    """
    # 创建文档
    doc = Document()
    #doc.add_heading("Research Error Fields", level=1)

    # 创建表格，2列：Study 和 Error Fields
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # 填充表头
    table.cell(0, 0).text = "Study"
    table.cell(0, 1).text = "Error Fields"

    # 填充数据
    for item in json_data:
        row_cells = table.add_row().cells
        row_cells[0].text = item["study"]
        row_cells[1].text = item["error_fields"]

    # 保存文档
    doc.save(output_file)
    print(f"Word document saved as {output_file}")


# 示例使用
json_data = [
  {
    "study": "Kraus et al. [20]",
    "error_fields": "Experimental group ethnicity (%),Control group ethnicity (%),Pharmacogenomic group results (SD),Treatment as usual group results (SD),GRADE"
  },
  {
    "study": "Mosley et al. [23]",
    "error_fields": "Experimental group male (%),Experimental group female (%),Control group male (%),Control group female (%),Experimental group ethnicity (%),Control group ethnicity (%),Opioids prescribed,GRADE"
  },
  {
    "study": "Agullo et al. [19]",
    "error_fields": "Control group mean age in years (SD),Experimental group ethnicity (%),Control group ethnicity (%),Pharmacogenomic group results (SD),Treatment as usual group results (SD),GRADE,P value"
  },
  {
    "study": "Hamilton et al. [21]",
    "error_fields": "Pharmacogenomic test used,Experimental group male (%),Experimental group female (%),Control group male (%),Control group female (%),Experimental group ethnicity (%),Control group ethnicity (%),Pharmacogenomic group results (SD),Treatment as usual group results (SD),GRADE,P value"
  },
  {
    "study": "Thomas et al. [24]",
    "error_fields": "Pharmacogenomic test used,Experimental group mean age in years (SD),Control group mean age in years (SD),Experimental group ethnicity (%),Control group ethnicity (%),Pharmacogenomic group results (SD),Treatment as usual group results (SD),GRADE,P value"
  },
  {
    "study": "Hamilton et al. [22]",
    "error_fields": "Sample size (patients),Pharmacogenomic test used,Experimental group male (%),Experimental group female (%),Experimental group mean age in years (SD),Control group male (%),Control group female (%),Control group mean age in years (SD),Experimental group ethnicity (%),Control group ethnicity (%),Opioids prescribed,Prescribing guidelines used,Pharmacogenomic group results (SD),Treatment as usual group results (SD),GRADE,P value"
  }
]


json_to_docx_table(json_data, "D:\\project\\zky\\paperAgent\\result\\research_errors.docx")
