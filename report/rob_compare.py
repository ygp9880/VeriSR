from typing import List, Dict
import os;
from utils import content_utils
import json
from docx import Document
from openai import  OpenAI
from utils.client_utils import get_client
from extract.extract_info import indify_rob2_table;
from extract.extract_info import match_study_file;
client = get_client();

def generate_rob2_summary_prompt(rob2_data_variable_name: str = "rob2_data") -> str:
    """
    Generate a prompt for a large language model to summarize ROB2 assessment results
    in no more than two sentences.

    Parameters:
    - rob2_data_variable_name: the variable name in your context that holds the ROB2 data

    Returns:
    - A string containing the prompt.
    """
    prompt = (
        f"Summarize the overall findings from the {rob2_data_variable_name} ROB2 assessment results "
        "in no more than two sentences, highlighting key patterns of agreement and disagreement across studies and domains."
    )
    return prompt

def json_to_docx_multilevel_table(json_str, output_file="output.docx"):
    """
    将 JSON 转换为 Word 文档多层表格：
    - 每个 study 一行
    - 每个 domain 占一组三列：extracted / computed / match
    """
    data = json.loads(json_str)
    if not data:
        return

    # 获取 domain 列名
    domains = list(data[0]["domains"].keys())

    # 创建 Word 文档
    doc = Document()
    doc.add_heading("Study Domains Table", level=1)

    # 计算列数：Study + 3列 * domain数
    num_cols = 1 + len(domains) * 3
    table = doc.add_table(rows=2+len(data), cols=num_cols)
    table.style = "Table Grid"

    # 第一行：大列
    table.cell(0,0).text = "Study"
    for j, domain in enumerate(domains):
        cell = table.cell(0, j*3+1)
        cell.text = domain
        # 合并大列的三列
        table.cell(0, j*3+1).merge(table.cell(0, j*3+3))

    # 第二行：子列
    for j in range(len(domains)):
        table.cell(1, j*3+1).text = "SR.risk"
        table.cell(1, j*3+2).text = "Model risk"
        table.cell(1, j*3+3).text = "Result"

    # 填充每个 study 的数据
    for i, study in enumerate(data, start=2):
        table.cell(i, 0).text = study["study"]
        for j, domain in enumerate(domains):
            d = study["domains"][domain]
            table.cell(i, j*3+1).text = str(d["extracted"])
            table.cell(i, j*3+2).text = str(d["computed"])
            table.cell(i, j*3+3).text = str(d["match"])

    # 保存 Word
    doc.save(output_file)
    print(f"Word 文档已保存为 {output_file}")


def compare_rob2(extracted: Dict, computed: Dict) -> List[Dict]:
    """
    Compare ROB2 results from paper extraction vs computed by framework.

    Parameters
    ----------
    extracted : dict
        Extracted ROB2 results from the article (simplified table with +/X/-)
    computed : dict
        Computed ROB2 JSON results from rob2 framework

    Returns
    -------
    List[Dict]
        A list of dictionaries with comparison results for each study and domain
    """
    # Mapping rob2 judgement to simplified symbols
    judgement_map = {
        "Low": "+",
        "Some concerns": "X",
        "High": "-"
    }

    results = []

    # Build a lookup for computed results
    computed_lookup = {}
    for study_json in computed.get("studies", []):
        study_name = study_json["study"]
        domain_map = {}
        for domain in study_json["domains"]:
            # Take the overall judgement for each domain
            domain_map[domain["index"]] = judgement_map.get(domain["judgment"], "?")
        # Overall
        domain_map["Overall"] = judgement_map.get(study_json.get("overall", "?"), "?")
        computed_lookup[study_name] = domain_map

    # Compare each row in extracted
    for row in extracted["rows"]:
        study_name = row[0]
        comparison = {"study": study_name, "domains": {}}
        if study_name not in computed_lookup:
            comparison["note"] = "Study not found in computed results"
        else:
            comp_domains = computed_lookup[study_name]
            for i, domain_symbol in enumerate(row[1:], 1):
                key = extracted["headers"][i]
                computed_symbol = comp_domains.get(i if i <= 5 else "Overall", "?")
                comparison["domains"][key] = {
                    "extracted": domain_symbol,
                    "computed": computed_symbol,
                    "match": domain_symbol == computed_symbol
                }
        results.append(comparison)

    return results

def compare(meta_path, meta_name, data_path, save_path):
    meta_content = content_utils.read_content(meta_path);
    meta_json = json.loads(meta_content);
    tables = meta_json['tables'];
    rbo2_json = indify_rob2_table(tables);
    if not 'meta_analysis_tables' in rbo2_json:
        return;

    table_indexd = rbo2_json['meta_analysis_tables'];
    size = len(table_indexd);
    if size == 0:
        return;

    index = rbo2_json['meta_analysis_tables'][0]['index'];
    extracted_rob2 = tables[index - 1];

    rows = extracted_rob2['rows'];
    rob2_studies = [];
    for row in rows:
        rob2_studies.append(row[0])

     # extracted_rob2
    path = data_path;
    files = os.listdir(path);

    match_result = match_study_file(rob2_studies, files);
    studies = [];
    for file in files:
        rob2_file = path + "\\" + file;
        # content_utils.read_content(rob2_file);
        # print(f" rob2_file is {rob2_file}");
        content = content_utils.read_content(rob2_file);
        content_json = json.loads(content);
        rob2_obj = {};
        rob2_obj['overall'] = content_json['overall'];
        domains = content_json['domains'];
        judge_domains = [];
        index = 1;
        for domain in domains:
            judgement = domain['judgment'];
            data = {"index": index, "judgment": judgement};

            for key in domain:
                if key != 'judgment':
                    data[key] = domain[key];

            # print(f" judgement is {judgement}");
            index = index + 1;
            # print(f" data is {data}");
            judge_domains.append(data);

        rob2_obj['domains'] = judge_domains;
        study = match_result[file];

        rob2_obj['study'] = study;
        studies.append(rob2_obj);

    computed_rob2 = {"studies": studies};
    judgement_map = {
        "+": "Low",
        "X": "Some concerns",
        "-": "High"
    }


    comparison_results = compare_rob2(extracted_rob2, computed_rob2)
    json_str = json.dumps(comparison_results, ensure_ascii=False, indent=4)
    # print(json_str)
    save_file_path = f"{save_path}\\{meta_name}_rob2_compare.txt";
    content_utils.write_str_to_file(save_file_path, json_str);
    prompt = generate_rob2_summary_prompt(json_str);

    response = client.chat.completions.create(
        model="gpt-5",
        messages=[
            {"role": "system", "content": "你是一个擅长从医学学术论文中提取结构化信息的助手"},
            {"role": "user", "content": prompt}
        ]
    )

    response_text = response.choices[0].message.content
    response_text = response_text.replace("```json", "")
    response_text = response_text.replace("```", "");
    print(f" reponse is {response_text}");




#compare("D:\\project\\zky\\paperAgent\\all_txt\\SR1.txt", "SR1","D:\\project\\zky\\paperAgent\\all_txt\\SR1\\rob2_result","D:\\project\\zky\\paperAgent\\all_txt");

'''
meta_path = "D:\\project\\zky\\paperAgent\\all_txt\\SR1.txt";
meta_content = content_utils.read_content(meta_path);
meta_json = json.loads(meta_content);
tables = meta_json['tables'];

rows = tables[4]['rows'];
studies = [];
for row in rows:
    studies.append(row[0])
print(f" studies is {studies} ");

rob2_path = "D:\\project\\zky\\paperAgent\\all_txt\\SR1\\rob2_result";
files = os.listdir(rob2_path);
match_result = match_study_file(studies, files);
print(f" match_result is {match_result} ");
study_name = match_result['SR1Krausetal2023.txt'];
print(f" study_name is {study_name} ");
'''
# 示例调用
'''
extracted_rob2 = {
    "title": "Table 5. Risk of bias assessment.",
    "headers": ["Study", "D1", "D2", "D3", "D4", "D5", "Overall"],
    "rows": [
        ["Kraus et al", "+", "+", "+", "+", "+", "+"],
        ["Mosley et al", "+", "X", "-", "+", "+", "X"],
        ["Agullo et al", "+", "+", "+", "+", "+", "+"],
        ["Hamilton et al (2022)", "+", "+", "+", "+", "+", "+"],
        ["Thomas et al", "+", "+", "+", "-", "+", "-"],
        ["Hamilton et al (2020)", "+", "+", "+", "+", "+", "+"]
    ]
}
'''

'''
computed_rob2 = {
    "studies": [
        {
            "study": "Kraus et al",
            "overall_judgement": "High",
            "domains": [
                {"index": 1, "judgement": "High"},
                {"index": 2, "judgement": "High"},
                {"index": 3, "judgement": "High"},
                {"index": 4, "judgement": "Low"},
                {"index": 5, "judgement": "Some concerns"}
            ]
        },
        # ...其他研究
    ]
}
'''

# 对比

'''
comparison_results = compare_rob2(extracted_rob2, computed_rob2)
for res in comparison_results:
    print(res)
'''
#data_path


#print(f" prompt is {prompt} ");

file_path = "D:\\project\\zky\\paperAgent\\all_txt\\SR1_rob2_compare.txt";

#json_str = content_utils.read_content(file_path);

#json_to_docx_multilevel_table(json_str,"D:\\project\\zky\\paperAgent\\report1\\rob2_result.docx")
