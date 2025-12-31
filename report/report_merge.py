from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.content_utils import read_content;
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import os

# 1 report merge direct
#2.Data Extraction Audit  D:\project\zky\paperAgent\all_txt\SR1_table_2.txt
#generate_prompt_for_check_summary check table
# 3.4 3.5 summary 使用导向生成
# 4.4 根据 META分析结果得到 summary 画图

#2.4

#3.1 generate_error_extraction_prompt_en
def generate_prompt_for_check_summary(header, data, table):
    """
    Generate the English prompt used to summarize discrepancies from the `check` field
    across included studies (Table 2, Table 3, Table 4).
    """

    prompt = f"""
You are an expert in systematic review quality verification and data extraction auditing.

Your task is to analyze the provided JSON data, especially the contents of the **"check"** field for each included study, and generate a structured discrepancy summary.

header is {header},
data is {data}
table ttile is {table}

## Dynamic Table Handling
- The dataset may contain discrepancies related to *Table 1, Table 2, Table 3, Table 4,* or other tables.
- **Do NOT assume fixed table numbers.**
- Instead, detect which tables are referenced in the "check" text and generate corresponding sections.
- For every table name that appears (e.g., “Table 2”, “Table 3”, “Table 4”), create a section titled:

### Table X – Main Issues

## Output Requirements
For each detected Table:
- Identify and group recurring discrepancy types into **bullet-point categories**.
- Each issue category must include:
  1. **A short category title**
  2. **A concise explanation**
  3. **Specific study examples** extracted from the `check` field  
     (e.g., “Hamilton et al. [21]”)

## Examples of Possible Issue Categories (not exhaustive):
- Gene panel mismatch or incomplete reporting  
- Baseline demographic inconsistencies  
- Errors in outcome data type or statistical reporting  
- Missing extraction of key variables  
- Incorrect sample size or participant flow  
- Misalignment between reported dataset and actual analysis population  

## Content Rules
- Only use information found in the `check` field of each study.
- Summaries should focus on **issue types**, not rewrite long texts.
- Group similar discrepancies across studies under the same category.
- Maintain an analytical, concise tone.

## Final Output Format Example (structure only)

### Table X – Main Issues
  1.Issue category 1
    brief explanation

  2. Issue category 2
     explanation 
- ...


"""

    return prompt

def json_to_word(json_data, output_filename='output.docx'):
    """
    将JSON数据转换为格式化的Word文档

    参数:
        json_data: JSON字符串或字典
        output_filename: 输出的Word文件名
    """
    # 如果输入是字符串，解析为字典
    if isinstance(json_data, str):
        data = json.loads(json_data)
    else:
        data = json_data

    # 创建Word文档
    doc = Document()

    # 添加标题
    p2 = doc.add_paragraph();
    p2.add_run("Statement: This report is an automated, methodology-focused summary based on the original text. It is intended to provide professionals with a rapid, preliminary assessment and does not substitute for a full expert manual review. Nothing in this report constitutes clinical decision advice. Users bear sole responsibility for any decisions made on the basis of this report.");


    doc.add_heading('一. Review Introduction', level=1)

    #doc.add_heading('Title', level=2)

    p = doc.add_paragraph()
    p.add_run("Title: ").bold = True
    p.add_run(str(data.get('title', 'N/A')))
    p.paragraph_format.space_after = Pt(12)

    # 添加目标/目的
    p = doc.add_paragraph()
    p.add_run("Objective: ").bold = True
    p.add_run(str(data.get('objective', 'N/A')))
    p.paragraph_format.space_after = Pt(12)


    # 添加研究信息
    #doc.add_heading('Sample Size', level=2)
    p = doc.add_paragraph()
    p.add_run("Sample Size: ").bold = True
    p.add_run(str(data.get('sample_size', 'N/A')))
    p.paragraph_format.space_after = Pt(12)


    # 添加纳入标准
    doc.add_paragraph('Inclusion Criteria', style='Heading 1')
    inclusion = data.get('inclusion_exclusion_criteria', {})
    for key, value in inclusion.items():
        p = doc.add_paragraph()
        p.add_run(f'{key}: ').bold = True
        p.add_run(str(value))
        p.paragraph_format.left_indent = Inches(0.25)

    # 添加主要结果
    doc.add_paragraph('Results', style='Heading 1')
    main_results = data.get('results', [])
    for i, result in enumerate(main_results, 1):
        p = doc.add_paragraph(result, style='List Number')
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('二. Report Body', level=1)
    # 保存文档
    doc.save(output_filename)
    print(f'Word文档已成功创建: {output_filename}')

    return output_filename


def add_text_to_doc(doc: Document, text: str, font_name='Times New Roman', font_size=12):
    """
    将文本添加到已有的 docx Document 对象中（不保存文件）。

    参数：
    doc (Document): 已创建或打开的 Document 对象
    text (str): 要添加的文本
    font_name (str): 字体名称，默认 'Times New Roman'
    font_size (int): 字体大小，默认 12
    """
    paragraph = doc.add_paragraph(text)
    run = paragraph.runs[0]
    run.font.name = font_name
    run.font.size = Pt(font_size)
    return doc


def extract_error_domains(data, sep=", "):
    """
    提取 match == False 的 domain 名称，并拼接成字符串
    """
    result = []

    for item in data:
        study = item.get("study")
        domains = item.get("domains", {})

        error_domain_list = [
            domain
            for domain, values in domains.items()
            if values.get("match") is False
        ]

        if error_domain_list:
            result.append({
                "study": study,
                "error_domains": sep.join(error_domain_list)
            })

    return result


def error_json_to_docx_table(error_json, doc):
    """
    将 Error Domain 汇总 JSON 写入 docx 表格
    不执行 save
    """
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Study"
    hdr_cells[1].text = "Error Domains"

    for item in error_json:
        row_cells = table.add_row().cells
        row_cells[0].text = item.get("study", "")
        row_cells[1].text = item.get("error_domains", "")

    return doc


def extract_data_check_errors(meta_json):
    """
    从 meta-analysis JSON 中提取 data_check 为 False 的 study

    返回格式：
    [
        {"study": "Thomas et al 2021", "Error Type": "Data Wrong"},
        ...
    ]
    """
    error_list = []

    for fig in meta_json:
        studies = fig.get("studies", [])
        for s in studies:
            if s.get("data_check") is False:
                error_list.append({
                    "study": s.get("study"),
                    "Error Type": "Data Wrong"
                })

    return error_list


def error_list_to_docx_table(error_list, doc):
    """
    将错误列表写入 docx 表格

    error_list: extract_data_check_errors 的返回结果
    doc: python-docx 的 Document 对象
    """
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Study"
    hdr_cells[1].text = "Error Type"

    # 数据行
    for item in error_list:
        row_cells = table.add_row().cells
        row_cells[0].text = item["study"]
        row_cells[1].text = item["Error Type"]

    return table

def check_data(base_path, meta_name):
    data_files = os.listdir(base_path);

    file_paths = [];
    for file in data_files:
        if file.__contains__("txt") and file.__contains__(meta_name) and file.__contains__("table"):
            full_path = base_path + "\\" + file;
            file_paths.append(full_path);

    return file_paths;

def filter_imgs(base_path, meta_name):
    data_files = os.listdir(base_path);

    img_paths = [];
    for file in data_files:
        if file.__contains__("png") and file.__contains__(meta_name):
            full_path = base_path + "\\" + file;
            img_paths.append(full_path);

    return img_paths;

def add_image_to_document(document, image_path, width=None, height=None):
    """
    向 python-docx 的 Document 对象中添加图片

    参数：
    - document: docx.Document 对象
    - image_path: 图片路径（str）
    - width: 图片宽度（Inches 或 Cm），可选
    - height: 图片高度（Inches 或 Cm），可选
    """

    if width and height:
        document.add_picture(image_path, width=width, height=height)
    elif width:
        document.add_picture(image_path, width=width)
    elif height:
        document.add_picture(image_path, height=height)
    else:
        document.add_picture(image_path)

def json_files_to_word(json_files, doc):
    """
    将多个 JSON 文件生成 Word 表格，写入同一个文档。

    参数：
        json_files (list of str): JSON 文件路径列表
        output_docx (str): 输出 Word 文件路径
    """
    #doc = Document()

    for json_file in json_files:
        # 读取 JSON 文件
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)

        title = data.get('title', 'Table')
        headers = data.get('headers', [])
        table_data = data.get('data', [])

        # 添加标题
        doc.add_heading(title, level=2)

        if not headers or not table_data:
            continue  # 空表格跳过

        # 创建表格
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # 添加表头
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        # 添加数据行
        for row in table_data:
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)

        # 添加空行分隔不同表格
        doc.add_paragraph('')


def merger(base_path,meta_name,save_file_path):
    report1_file_path = f"{base_path}\\{meta_name}_report_1.json";


    with open(report1_file_path, 'r', encoding='utf-8') as f:
        json_data = json.load(f)

    json_to_word(json_data, save_file_path)

    title_map = {};
    title_map[1] = "1. Literature Search Audit";
    title_map[2] = "2. Data Extraction Audit";
    title_map[3] = "3. Risk of Bias Audit";
    title_map[4] = "4. Meta-Analysis Audit";
    title_map[5] = "5. Additional Analyses and Evidence Quality Assessment";

    doc = Document(save_file_path);

    for j in range(1, 6):
        report2_file_path = f"{base_path}\\{meta_name}_report_2_{j}.json";
        with open(report2_file_path, 'r', encoding='utf-8') as f:
            json_2_data = json.load(f)

        keys = json_2_data.keys();

        index = j;
        title = title_map[index];
        print(f" title is {title}");
        doc.add_heading(title, level=2)
        count = 1;
        for key in keys:
            if j == 2 and count == 3:
                small_title = "2.3 Content Verification";
                heading = doc.add_heading(small_title, level=3)
                run = heading.runs[0]  # 获取标题的第一个 run
                run.font.size = Pt(10)

                report2_file_path = f"{base_path}\\{meta_name}_report_2_2_3.txt";
                report2_content = read_content(report2_file_path);
                add_text_to_doc(doc, report2_content);

                file_paths = check_data(base_path, meta_name);
                json_files_to_word(file_paths, doc);
                count = count + 1;

            if j == 4 and count == 3:
                small_title = "4.4 Content Verification";
                heading = doc.add_heading(small_title, level=3)
                run = heading.runs[0]  # 获取标题的第一个 run
                run.font.size = Pt(10)

                report2_file_path = f"{base_path}\\{meta_name}_report_2_4_4.txt";
                report2_content = read_content(report2_file_path);
                add_text_to_doc(doc, report2_content);
                imgs_paths = filter_imgs(base_path, meta_name);
                for img_path in imgs_paths:
                    add_image_to_document(doc, img_path,  width=Inches(7));
                #add_image_to_document(document=doc, image_path="example.png", width=Inches(4))



            key_str: str = key;
            value = json_2_data[key];
            process_key_str = key_str.replace("_", " ");
            small_title = f"{index}.{count} {process_key_str} ";
            heading = doc.add_heading(small_title, level=3)
            run = heading.runs[0]  # 获取标题的第一个 run
            run.font.size = Pt(10)
            if isinstance(value, str):
                p = doc.add_paragraph(value)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = 1.15
                print(f" value is {value}");
            else:
                for key, item_value in value.items():
                    if key == "evidence":
                        continue
                    item = doc.add_paragraph(item_value)
                    item.style.font.size = Pt(10)

                if "evidence" in value:
                    # evidence 内容
                    # 分隔行
                    # 添加 Evidence 提示，斜体加粗
                    evidence_title = doc.add_paragraph("Evidence / Source:")
                    run_title = evidence_title.runs[0]
                    run_title.bold = True
                    run_title.italic = True
                    run_title.font.size = Pt(8)

                    # evidence 内容，字体小一点，显示为引用风格
                    evidence_para = doc.add_paragraph(value["evidence"])
                    evidence_para.style.font.size = Pt(8)  # 小字体
                    evidence_para.paragraph_format.left_indent = Pt(18)  # 缩进
                    evidence_para.paragraph_format.space_before = Pt(2)
                    evidence_para.paragraph_format.space_after = Pt(4)
            count = count + 1;








    doc.add_heading('三. Report Summary', level=1)
    doc.add_heading("1. Overall Conclusion", level=2)

    report3_file_path = f"{base_path}\\{meta_name}_report_3_1.json";
    with open(report3_file_path, 'r', encoding='utf-8') as f:
        json3_data = json.load(f)

    keys = json3_data.keys();
    for key in keys:
        key_str: str = key;
        value = json3_data[key];
        process_key_str = key_str.replace("_", " ");
        small_title = process_key_str;
        heading = doc.add_heading(small_title, level=3)
        run = heading.runs[0]  # 获取标题的第一个 run
        run.font.size = Pt(10)
        if isinstance(value, str):
            p = doc.add_paragraph(value)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.15
            print(f" value is {value}");

    doc.add_heading("1. Items Requiring Focused", level=2)
    heading = doc.add_heading("Data Extraction", level=3)
    run = heading.runs[0]  # 获取标题的第一个 run
    run.font.size = Pt(10)

    # report_3_2_1

    report3_2_1_file_path = f"{base_path}\\{meta_name}_report_3_2_1.json";
    with open(report3_2_1_file_path, 'r', encoding='utf-8') as f:
        json3_2_1_data = json.load(f)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # 填充表头
    table.cell(0, 0).text = "Study"
    table.cell(0, 1).text = "Error Fields"

    # 填充数据
    for item in json3_2_1_data:
        row_cells = table.add_row().cells
        row_cells[0].text = item["study"]
        row_cells[1].text = item["error_fields"]



    heading = doc.add_heading("Meta-analysis", level=3)
    run = heading.runs[0]  # 获取标题的第一个 run
    run.font.size = Pt(10)

    #extract_result_all_SR1.json

    result_file_path = f"{base_path}\\extract_result_all_{meta_name}.json";

    content = read_content(result_file_path);
    json_data = json.loads(content);
    result = extract_data_check_errors(json_data);

    error_list_to_docx_table(result, doc);


    heading = doc.add_heading("Risk of bais assessment", level=3)
    run = heading.runs[0]  # 获取标题的第一个 run
    run.font.size = Pt(10)

    #SR1_rob2_compare.txt
    file_path = f"{base_path}\\{meta_name}_rob2_compare.txt";
    content = read_content(file_path);
    json_data = json.loads(content);
    result = extract_error_domains(json_data);

    error_json_to_docx_table(result,doc);


    doc.add_heading('四. Attachment', level=1)


    data_files = os.listdir(base_path);

    for file in data_files:
        file_path = f"{base_path}\\{file}";
        if file.__contains__("txt") and file.__contains__(meta_name) and file.__contains__("rob2") and (not file.__contains__("compare")):
            txt = read_content(file_path);
            add_text_to_doc(doc, txt);


    doc.save(save_file_path)




#merger("D:\\project\\zky\\paperAgent\\all_txt","SR1","D:\\project\\zky\\paperAgent\\report_doc\\SR1_report.docx");
