import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.content_utils import read_content;
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
def json_to_word(json_data, output_filename='output.docx'):
    """
    将JSON数据转换为格式化的Word文档

    参数:
        json_data: JSON字符串或字典
        output_filename: 输出的Word文件名
    """
    # 如果输入是字符串，解析为字典
    if isinstance(json_data, str):
        data = json.loads(json_data)
    else:
        data = json_data

    # 创建Word文档
    doc = Document()

    # 添加标题
    p2 = doc.add_paragraph();
    p2.add_run("Statement: This report is an automated, methodology-focused summary based on the original text. It is intended to provide professionals with a rapid, preliminary assessment and does not substitute for a full expert manual review. Nothing in this report constitutes clinical decision advice. Users bear sole responsibility for any decisions made on the basis of this report.");


    doc.add_heading('一. Review Introduction', level=1)

    #doc.add_heading('Title', level=2)

    p = doc.add_paragraph()
    p.add_run("Title: ").bold = True
    p.add_run(str(data.get('title', 'N/A')))
    p.paragraph_format.space_after = Pt(12)

    # 添加目标/目的
    p = doc.add_paragraph()
    p.add_run("Objective: ").bold = True
    p.add_run(str(data.get('objective', 'N/A')))
    p.paragraph_format.space_after = Pt(12)


    # 添加研究信息
    #doc.add_heading('Sample Size', level=2)
    p = doc.add_paragraph()
    p.add_run("Sample Size: ").bold = True
    p.add_run(str(data.get('sample_size', 'N/A')))
    p.paragraph_format.space_after = Pt(12)


    # 添加纳入标准
    doc.add_paragraph('Inclusion Criteria', style='Heading 1')
    inclusion = data.get('inclusion_exclusion_criteria', {})
    for key, value in inclusion.items():
        p = doc.add_paragraph()
        p.add_run(f'{key}: ').bold = True
        p.add_run(str(value))
        p.paragraph_format.left_indent = Inches(0.25)

    # 添加主要结果
    doc.add_paragraph('Results', style='Heading 1')
    main_results = data.get('results', [])
    for i, result in enumerate(main_results, 1):
        p = doc.add_paragraph(result, style='List Number')
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('二. Report Body', level=1)
    # 保存文档
    doc.save(output_filename)
    print(f'Word文档已成功创建: {output_filename}')

    return output_filename


def add_methodology_to_word(doc, json_data):
    """
    将方法学相关的JSON数据添加到现有Word文档

    参数:
        doc: Document对象（现有的Word文档）
        json_data: JSON字符串或字典
    """
    # 如果输入是字符串，解析为字典
    if isinstance(json_data, str):
        data = json.loads(json_data)
    else:
        data = json_data

    # 添加分页符
    doc.add_page_break()

    # 添加方法学部分的主标题
    doc.add_heading('二. Report Body', level=1)
    doc.add_heading('1. Literature Search Audit', level=1)

    # 1. 搜索策略
    if 'Search_Strategy' in data:
        doc.add_heading('Search Strategy', level=3)
        p = doc.add_paragraph(data['Search_Strategy'])
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15

    # 2. 筛选流程
    if 'Screening_Process' in data:
        doc.add_heading('Screening Process', level=3)
        p = doc.add_paragraph(data['Screening_Process'])
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15

    # 3. 结果呈现
    if 'Results_Presentation' in data:
        doc.add_heading('Results Presentation', level=3)
        p = doc.add_paragraph(data['Results_Presentation'])
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15

    # 4. 方法学总结
    if 'Methodology_Summary' in data:
        doc.add_heading('Methodology Summary', level=3)
        p = doc.add_paragraph(data['Methodology_Summary'])
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15

        # 为总结添加浅色背景效果（通过添加表格实现）
        # 注：python-docx不直接支持段落背景色，这里用边框来突出显示
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # 添加底纹
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F0F0F0')  # 浅灰色背景
        p._element.get_or_add_pPr().append(shading_elm)

    return doc



if __name__ == '__main__':
    # 读取JSON文件
    #report_1_json = read_content("D:\\project\\zky\\paperAgent\\result\\report_1.json");
    #json_to_word(report_1_json,"D:\\project\\zky\\paperAgent\\result\\output.docx");


    title_map = {};
    title_map[1] = "1. Literature Search Audit";
    title_map[2] = "2. Data Extraction Audit";
    title_map[3] = "3. Risk of Bias Audit";
    title_map[4] = "4. Meta-Analysis Audit";
    title_map[5] = "5. Additional Analyses and Evidence Quality Assessment";

    index = 5;

    with open(f'D:\\project\\zky\\paperAgent\\result\\report_3_{index}.json', 'r', encoding='utf-8') as f:
        json_data = json.load(f)

    # 转换为Word
    #json_to_word(json_data, 'output.docx')
    keys = json_data.keys();
    doc = Document('D:\\project\\zky\\paperAgent\\result\\output.docx');
    title = title_map[index];
    print(f" title is {title}");
    doc.add_heading(title, level=2)
    count = 1;
    for key in keys:
        key_str:str = key;
        value = json_data[key];
        process_key_str = key_str.replace("_", " ");
        small_title = f"{index}.{count} {process_key_str} ";
        heading = doc.add_heading(small_title, level=3)
        run = heading.runs[0]  # 获取标题的第一个 run
        run.font.size = Pt(10)
        if isinstance(value, str):
            p = doc.add_paragraph(value)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.15
            print(f" value is {value}");
        else:
            for key, item_value in value.items():
                if key == "evidence":
                    continue
                item = doc.add_paragraph(item_value)
                item.style.font.size = Pt(10)

            if "evidence" in value:
                # evidence 内容
                # 分隔行
                # 添加 Evidence 提示，斜体加粗
                evidence_title = doc.add_paragraph("Evidence / Source:")
                run_title = evidence_title.runs[0]
                run_title.bold = True
                run_title.italic = True
                run_title.font.size = Pt(8)

                # evidence 内容，字体小一点，显示为引用风格
                evidence_para = doc.add_paragraph(value["evidence"])
                evidence_para.style.font.size = Pt(8)  # 小字体
                evidence_para.paragraph_format.left_indent = Pt(18)  # 缩进
                evidence_para.paragraph_format.space_before = Pt(2)
                evidence_para.paragraph_format.space_after = Pt(4)


        count = count + 1;
    #print(f" json_data is {json_data}");
    doc.save("D:\\project\\zky\\paperAgent\\result\\output.docx")

